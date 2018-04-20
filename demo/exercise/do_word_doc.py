#!/usr/bin/env python3
# encoding:utf-8
'''
@author: lierl
@file: do_word_doc.py
@time: 2018/4/17 20:21
'''
__author__ = 'lierl'

import docx

def read_docx(file_name):
    doc = docx.Document(file_name)
    content = '\n'.join([para.text for para in doc.paragraphs])
    return content

def read_tables(file_name):
    doc = docx.Document(file_name)
    for table in doc.tables:  # 遍历所有表格
        print("------------------------------------")
        for row in table.rows:  # 遍历表格的所有行
            for cell in row.cells:
                print(cell.text, end='、')
            print()
print(read_tables("E:\\aa.docx"))
# print(read_docx("E:\\aa.docx"))
